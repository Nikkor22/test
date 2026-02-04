import os
import io
from datetime import datetime
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.services.gpt_service import GPTService


WORK_TYPE_NAMES = {
    "homework": "Домашняя работа",
    "lab": "Лабораторная работа",
    "practical": "Практическая работа",
    "coursework": "Курсовая работа",
    "report": "Реферат",
    "essay": "Эссе",
    "presentation": "Презентация",
    "exam": "Экзамен",
    "test": "Контрольная работа",
}


class WorkGeneratorService:
    def __init__(self, gpt_service: GPTService):
        self.gpt_service = gpt_service

    async def generate_work_content(
        self,
        subject_name: str,
        work_type: str,
        work_number: Optional[int],
        title: str,
        description: Optional[str],
        materials: List[str]
    ) -> str:
        """Генерирует содержимое работы на основе материалов предмета"""

        work_type_name = WORK_TYPE_NAMES.get(work_type, work_type)
        work_title = f"{work_type_name}"
        if work_number:
            work_title += f" №{work_number}"
        work_title += f": {title}"

        # Prepare context from materials
        materials_context = ""
        if materials:
            materials_context = "\n\n".join(materials[:5])  # Limit to 5 materials
            # Truncate if too long
            if len(materials_context) > 15000:
                materials_context = materials_context[:15000] + "\n...[материалы сокращены]"

        prompt = f"""
Предмет: {subject_name}
Тип работы: {work_type_name}
Номер работы: {work_number or "не указан"}
Название: {title}
Описание задания: {description or "не указано"}

Материалы по предмету:
{materials_context if materials_context else "Материалы не загружены"}

Задача: Напиши полный текст работы "{work_title}" по предмету "{subject_name}".

Требования к работе:
1. Работа должна быть структурированной (введение, основная часть, заключение)
2. Используй материалы предмета для наполнения контентом
3. Объем работы должен быть достаточным для типа работы (лабораторная - 3-5 страниц, реферат - 10-15 страниц, и т.д.)
4. Пиши академическим языком, но понятно
5. Добавь примеры и пояснения где уместно
6. Если есть формулы или расчеты - приведи их
7. В конце добавь список использованных источников (если применимо)

НЕ ДОБАВЛЯЙ титульный лист - он будет добавлен отдельно.
Начни сразу с содержания работы.
"""

        try:
            response = await self.gpt_service.client.chat.completions.create(
                model="openai/gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": """Ты профессиональный помощник для написания учебных работ.
Твоя задача - создавать качественные, структурированные и содержательные работы на основе предоставленных материалов.
Пиши грамотно, академическим стилем. Используй подзаголовки для структурирования текста.
Объем работы должен соответствовать её типу."""
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=4000
            )

            return response.choices[0].message.content
        except Exception as e:
            print(f"Work generation error: {e}")
            raise Exception(f"Не удалось сгенерировать работу: {str(e)}")

    async def create_document(
        self,
        content: str,
        subject_name: str,
        work_type: str,
        work_number: Optional[int],
        student_name: str,
        group_number: str,
        template_path: Optional[str],
        output_dir: str,
        user_id: int,
        deadline_id: int
    ) -> Tuple[str, str]:
        """Создает DOCX документ с титульным листом и содержимым"""

        work_type_name = WORK_TYPE_NAMES.get(work_type, work_type)
        current_date = datetime.now().strftime("%d.%m.%Y")

        # Create document
        if template_path and os.path.exists(template_path):
            # Use template
            doc = Document(template_path)
            # Replace placeholders in template
            self._replace_placeholders(
                doc,
                subject_name=subject_name,
                work_type=work_type_name,
                work_number=work_number,
                student_name=student_name,
                group_number=group_number,
                date=current_date
            )
        else:
            # Create default title page
            doc = Document()
            self._create_default_title_page(
                doc,
                subject_name=subject_name,
                work_type=work_type_name,
                work_number=work_number,
                student_name=student_name,
                group_number=group_number,
                date=current_date
            )

        # Add page break after title
        doc.add_page_break()

        # Add content
        self._add_content(doc, content)

        # Save document
        work_number_str = f"_{work_number}" if work_number else ""
        file_name = f"{work_type}{work_number_str}_{subject_name.replace(' ', '_')}_{deadline_id}.docx"
        file_name = "".join(c for c in file_name if c.isalnum() or c in ('_', '-', '.'))  # Sanitize
        file_path = os.path.join(output_dir, f"{user_id}_{file_name}")

        doc.save(file_path)

        return file_name, file_path

    def _replace_placeholders(
        self,
        doc: Document,
        subject_name: str,
        work_type: str,
        work_number: Optional[int],
        student_name: str,
        group_number: str,
        date: str
    ):
        """Заменяет плейсхолдеры в шаблоне"""
        replacements = {
            "{{subject_name}}": subject_name,
            "{{date}}": date,
            "{{work_type}}": work_type,
            "{{work_number}}": str(work_number) if work_number else "",
            "{{student_name}}": student_name,
            "{{group_number}}": group_number,
            # Also support Russian placeholders
            "{{предмет}}": subject_name,
            "{{дата}}": date,
            "{{тип_работы}}": work_type,
            "{{номер_работы}}": str(work_number) if work_number else "",
            "{{имя_студента}}": student_name,
            "{{группа}}": group_number,
        }

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

    def _create_default_title_page(
        self,
        doc: Document,
        subject_name: str,
        work_type: str,
        work_number: Optional[int],
        student_name: str,
        group_number: str,
        date: str
    ):
        """Создает титульный лист по умолчанию"""
        # University header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ\nРОССИЙСКОЙ ФЕДЕРАЦИИ")
        run.font.size = Pt(12)
        run.bold = True

        # Empty lines
        for _ in range(3):
            doc.add_paragraph()

        # Work type and number
        work_title = doc.add_paragraph()
        work_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = work_type.upper()
        if work_number:
            title_text += f" №{work_number}"
        run = work_title.add_run(title_text)
        run.font.size = Pt(16)
        run.bold = True

        # Subject
        subject_para = doc.add_paragraph()
        subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subject_para.add_run(f"по дисциплине «{subject_name}»")
        run.font.size = Pt(14)

        # Empty lines
        for _ in range(5):
            doc.add_paragraph()

        # Student info (right-aligned)
        student_info = doc.add_paragraph()
        student_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = student_info.add_run(f"Выполнил: {student_name}")
        run.font.size = Pt(12)

        if group_number:
            group_para = doc.add_paragraph()
            group_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = group_para.add_run(f"Группа: {group_number}")
            run.font.size = Pt(12)

        # Empty lines
        for _ in range(5):
            doc.add_paragraph()

        # Date and city
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"Дата выполнения: {date}")
        run.font.size = Pt(12)

    def _add_content(self, doc: Document, content: str):
        """Добавляет содержимое работы в документ"""
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # Check if it's a heading (starts with # or all caps short line)
            if line.startswith('# '):
                # Main heading
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # Subheading
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Sub-subheading
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                para = doc.add_paragraph(line[2:], style='List Bullet')
            elif len(line) > 0 and line[0].isdigit() and '. ' in line[:4]:
                # Numbered list
                para = doc.add_paragraph(line, style='List Number')
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.line_spacing = 1.5
